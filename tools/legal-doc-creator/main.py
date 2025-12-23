"""Generate Contoso documentation in Markdown and convert to Word."""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document


COMPANY_NAME = "Contoso Corporation"


@dataclass
class DocumentTemplate:
	category: str
	name: str
	title: str
	filename: str
	summary: str
	body_sections: list[tuple[str, str]]


def build_templates() -> list[DocumentTemplate]:
	today = datetime.utcnow().date().isoformat()

	legal = [
		DocumentTemplate(
			category="Legal",
			name="code-of-conduct",
			title="Code of Conduct and Ethics",
			filename=f"legal-code-of-conduct-{today}.md",
			summary=(
				"Standards for ethical behavior, compliance, and reporting across "
				f"{COMPANY_NAME} manufacturing operations."
			),
			body_sections=[
				(
					"Purpose and Principles",
					"- Integrity in sourcing, production, and sales\n"
					"- Zero tolerance for bribery, forced labor, or unsafe work\n"
					"- Whistleblowing routes with anti-retaliation safeguards",
				),
				(
					"Roles and Accountability",
					"- Employees follow policy; managers reinforce and coach\n"
					"- Legal and Compliance owns governance and training\n"
					"- Suppliers sign annual attestations and submit audits",
				),
				(
					"Controls and Monitoring",
					"- Annual training with completion tracking\n"
					"- Third-party due diligence for high-risk regions\n"
					"- Quarterly risk reviews and incident root-cause analyses",
				),
			],
		),
		DocumentTemplate(
			category="Legal",
			name="supplier-agreement",
			title="Supplier Agreement Framework",
			filename=f"legal-supplier-agreement-{today}.md",
			summary=(
				"Standard terms for quality, delivery, ethics, and liability with "
				f"manufacturing suppliers serving {COMPANY_NAME}."
			),
			body_sections=[
				(
					"Scope and Quality",
					"- Product specs, change control, and PPAP alignment\n"
					"- Lot traceability and defect thresholds\n"
					"- Right-to-audit with notice provisions",
				),
				(
					"Commercial and Delivery",
					"- Incoterms, lead times, and safety stock expectations\n"
					"- Late-delivery credits and expedited shipping triggers\n"
					"- Currency, invoicing cadence, and tax handling",
				),
				(
					"Compliance and Liability",
					"- Anti-counterfeit and IP protection clauses\n"
					"- Data protection and confidentiality schedules\n"
					"- Indemnities, insurance limits, and termination rights",
				),
			],
		),
		DocumentTemplate(
			category="Legal",
			name="data-protection",
			title="Data Protection and Privacy Policy",
			filename=f"legal-data-protection-{today}.md",
			summary=(
				"Guardrails for handling personal and operational data across "
				f"Contoso systems and partner integrations."
			),
			body_sections=[
				(
					"Data Inventory and Classification",
					"- Map HR, supplier, and customer data flows\n"
					"- Classify data (Public, Internal, Confidential, Restricted)\n"
					"- Maintain records of processing activities",
				),
				(
					"Protection Measures",
					"- Encryption in transit and at rest\n"
					"- Least-privilege access and MFA\n"
					"- Retention schedules with defensible deletion",
				),
				(
					"Incident and Requests",
					"- Breach notification within regulatory windows\n"
					"- Data subject request handling playbook\n"
					"- Annual tabletop exercises and remediation tracking",
				),
			],
		),
	]

	travel = [
		DocumentTemplate(
			category="Travel",
			name="travel-policy",
			title="Corporate Travel Policy",
			filename=f"travel-policy-{today}.md",
			summary=(
				"Guidelines for booking, safety, approvals, and expenses for "
				f"Contoso travelers and on-site plant visits."
			),
			body_sections=[
				(
					"Planning and Booking",
					"- Book via approved agency; economy for under 5 hours\n"
					"- Preferred hotels with safety and proximity to sites\n"
					"- Visa and vaccination checklist before departure",
				),
				(
					"Safety and Conduct",
					"- Daily check-in protocol for high-risk regions\n"
					"- Emergency contacts and evacuation partners\n"
					"- Prohibit personal deviations without manager approval",
				),
				(
					"Expenses and Approvals",
					"- Per-diem guidance by country\n"
					"- Receipts required over local threshold\n"
					"- Submit reports within 7 days of return",
				),
			],
		),
		DocumentTemplate(
			category="Travel",
			name="travel-risk",
			title="Travel Risk Assessment and Emergency Protocol",
			filename=f"travel-risk-protocol-{today}.md",
			summary=(
				"Pre-trip risk controls and response actions for Contoso "
				"employees and contractors."
			),
			body_sections=[
				(
					"Risk Assessment",
					"- Country and city risk rating with advisory level\n"
					"- Site security review and transport plan\n"
					"- Health risks and required PPE for plant visits",
				),
				(
					"Preparedness",
					"- Traveler briefing and contact tree\n"
					"- Backup communication channels (satphone, SMS)\n"
					"- Medical support providers and clinic locations",
				),
				(
					"Incident Response",
					"- Escalation tiers and decision thresholds\n"
					"- Shelter-in-place vs. evacuation criteria\n"
					"- Post-incident debrief and corrective actions",
				),
			],
		),
	]

	hr = [
		DocumentTemplate(
			category="HR",
			name="employee-handbook",
			title="Employee Handbook Overview",
			filename=f"hr-employee-handbook-{today}.md",
			summary=(
				"Core employment practices, workplace standards, and employee "
				f"support programs at {COMPANY_NAME}."
			),
			body_sections=[
				(
					"Employment and Onboarding",
					"- Offer, probation, and background verification\n"
					"- Safety induction and role-based training\n"
					"- Code of conduct acknowledgement",
				),
				(
					"Compensation and Benefits",
					"- Pay calendar and allowance policies\n"
					"- Health coverage, EAP, and wellness programs\n"
					"- Leave types with approval flows",
				),
				(
					"Workplace Standards",
					"- Anti-discrimination and respectful workplace\n"
					"- Safety reporting and near-miss logging\n"
					"- Remote work and site access expectations",
				),
			],
		),
		DocumentTemplate(
			category="HR",
			name="performance-playbook",
			title="Performance Management Playbook",
			filename=f"hr-performance-playbook-{today}.md",
			summary=(
				"Goal setting, coaching, and review cadence for Contoso teams "
				"across plants and corporate offices."
			),
			body_sections=[
				(
					"Goal Setting",
					"- Annual objectives aligned to plant throughput and quality\n"
					"- Quarterly KPIs with leading indicators\n"
					"- Individual development plans tied to skills",
				),
				(
					"Coaching and Feedback",
					"- Monthly one-on-ones with documented actions\n"
					"- Peer feedback for cross-functional projects\n"
					"- Recognition for safety, quality, and cost savings",
				),
				(
					"Reviews and Remediation",
					"- Midyear and annual reviews with calibration\n"
					"- Performance improvement plans and timelines\n"
					"- Training or role reassignment options before exit",
				),
			],
		),
		DocumentTemplate(
			category="HR",
			name="hse-guidelines",
			title="Health, Safety, and Environment (HSE) Guidelines",
			filename=f"hr-hse-guidelines-{today}.md",
			summary=(
				"Safe manufacturing operations, environmental stewardship, and "
				"regulatory compliance requirements."
			),
			body_sections=[
				(
					"Safety Controls",
					"- PPE matrix by work area and task\n"
					"- Lockout-tagout and confined space procedures\n"
					"- Incident reporting and corrective action tracking",
				),
				(
					"Environment and Sustainability",
					"- Waste handling and recycling standards\n"
					"- Emissions monitoring and reporting\n"
					"- Spill response kits and training",
				),
				(
					"Training and Audits",
					"- Annual HSE curriculum with refreshers\n"
					"- Toolbox talks and pre-shift briefings\n"
					"- Internal audits and regulatory inspections",
				),
			],
		),
	]

	return legal + travel + hr


def render_markdown(template: DocumentTemplate) -> str:
	today = datetime.utcnow().date().isoformat()
	lines: list[str] = []
	lines.append(f"# {template.title}")
	lines.append(f"**Company:** {COMPANY_NAME}  ")
	lines.append(f"**Category:** {template.category}  ")
	lines.append(f"**Document:** {template.name}  ")
	lines.append(f"**Last Updated:** {today}")
	lines.append("")
	lines.append("## Overview")
	lines.append(template.summary)
	lines.append("")
	for heading, body in template.body_sections:
		lines.append(f"## {heading}")
		lines.extend(body.split("\n"))
		lines.append("")
	lines.append("## Approvals and Ownership")
	lines.append("- Document owner: <role>\n- Approver: <role>\n- Review cycle: Annual or upon regulation change")
	lines.append("")
	lines.append("## Change Log")
	lines.append("- <date>: <summary of change>\n- <date>: <summary of change>")
	return "\n".join(lines).strip() + "\n"


def write_markdown_files(output_dir: Path, templates: Iterable[DocumentTemplate]) -> list[Path]:
	output_dir.mkdir(parents=True, exist_ok=True)
	written: list[Path] = []
	for template in templates:
		path = output_dir / template.filename
		path.write_text(render_markdown(template), encoding="utf-8")
		written.append(path)
	return written


def strip_basic_markdown(text: str) -> str:
	return text.replace("**", "").replace("`", "").strip()


def convert_markdown_file(md_path: Path, docx_path: Path) -> None:
	doc = Document()
	lines = md_path.read_text(encoding="utf-8").splitlines()
	for line in lines:
		stripped = line.strip()
		if not stripped:
			continue
		if stripped.startswith("### "):
			doc.add_heading(strip_basic_markdown(stripped[4:]), level=3)
		elif stripped.startswith("## "):
			doc.add_heading(strip_basic_markdown(stripped[3:]), level=2)
		elif stripped.startswith("# "):
			doc.add_heading(strip_basic_markdown(stripped[2:]), level=1)
		elif stripped.startswith("- "):
			doc.add_paragraph(strip_basic_markdown(stripped[2:]), style="List Bullet")
		elif stripped.startswith("---"):
			doc.add_paragraph("\u2014")
		else:
			doc.add_paragraph(strip_basic_markdown(stripped))
	docx_path.parent.mkdir(parents=True, exist_ok=True)
	doc.save(docx_path)


def convert_directory(input_dir: Path, output_dir: Path) -> list[Path]:
	output_dir.mkdir(parents=True, exist_ok=True)
	generated: list[Path] = []
	for md_file in input_dir.glob("*.md"):
		target = output_dir / (md_file.stem + ".docx")
		convert_markdown_file(md_file, target)
		generated.append(target)
	return generated


def parse_args() -> argparse.Namespace:
	parser = argparse.ArgumentParser(description="Contoso documentation generator")
	subparsers = parser.add_subparsers(dest="command", required=True)

	gen = subparsers.add_parser("generate", help="Generate Markdown files")
	gen.add_argument(
		"--output",
		type=Path,
		default=Path("output/markdown"),
		help="Directory to write Markdown files",
	)

	conv = subparsers.add_parser("convert", help="Convert Markdown to Word")
	conv.add_argument(
		"--input",
		type=Path,
		default=Path("output/markdown"),
		help="Directory containing Markdown files",
	)
	conv.add_argument(
		"--output",
		type=Path,
		default=Path("output/word"),
		help="Directory to write Word documents",
	)

	return parser.parse_args()


def main() -> None:
	args = parse_args()
	if args.command == "generate":
		templates = build_templates()
		written = write_markdown_files(args.output, templates)
		print(f"Generated {len(written)} Markdown files in {args.output}")
	elif args.command == "convert":
		generated = convert_directory(args.input, args.output)
		print(f"Converted {len(generated)} files to Word in {args.output}")


if __name__ == "__main__":
	main()
